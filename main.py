import os
import json
from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import ApplicationBuilder, CommandHandler, CallbackQueryHandler, CallbackContext, ConversationHandler, MessageHandler, filters
import re
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

FULL_NAME_REGEX = r'^[А-ЯЇЄІҐ][а-яїєіґ]{2,} [А-ЯЇЄІҐ][а-яїєіґ]{2,} [А-ЯЇЄІҐ][а-яїєіґ]{2,}$'

load_dotenv()

# Conversation states
ROLE, FACULTY, DEPARTMENT, DOCUMENT, FULL_NAME, ADDITIONAL_DATA, CONFIRMATION = map(chr, range(7))
USER_DATA = {}

def load_data():
    try:
        with open('data_uk.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print("Помилка: Файл data_uk.json не знайдено.")
        return {}

data = load_data()
ui_text = data.get('ui_text', {})

def generate_document(user_data, output_format='docx'):
    file_name = f"{user_data['full_name']}_{user_data['document']}"
    if output_format.lower() == 'docx':
        doc = Document()
        doc.add_paragraph(f"Document: {user_data['document']}")
        doc.add_paragraph(f"Full Name: {user_data['full_name']}")
        if 'additional_data' in user_data:
            for key, value in user_data['additional_data'].items():
                doc.add_paragraph(f"{key.title()}: {value}")
        file_path = f"{file_name}.docx"
        doc.save(file_path)
        return file_path
    elif output_format.lower() == 'pdf':
        file_path = f"{file_name}.pdf"
        c = canvas.Canvas(file_path, pagesize=letter)
        width, height = letter
        y_position = height - 50
        c.drawString(100, y_position, f"Document: {user_data['document']}")
        y_position -= 20
        c.drawString(100, y_position, f"Full Name: {user_data['full_name']}")
        y_position -= 20
        if 'additional_data' in user_data:
            for key, value in user_data['additional_data'].items():
                c.drawString(100, y_position, f"{key.title()}: {value}")
                y_position -= 20
        c.showPage()
        c.save()
        return file_path
    else:
        raise ValueError("Непідтримуваний формат. Використовуйте 'docx' або 'pdf'.")

def get_current_selection(user_data):
    selection = ""
    if 'role' in user_data:
        selection += f"{ui_text.get('role_label', 'Роль')}: {user_data['role']}\n"
    if 'faculty' in user_data:
        selection += f"{ui_text.get('faculty_label', 'Факультет')}: {user_data['faculty']}\n"
    if 'department' in user_data:
        selection += f"{ui_text.get('department_label', 'Кафедра')}: {user_data['department']}\n"
    if 'document' in user_data:
        selection += f"{ui_text.get('document_label', 'Документ')}: {user_data['document']}\n"
    return selection

async def start(update: Update, context: CallbackContext) -> int:
    user_id = update.effective_user.id
    USER_DATA[user_id] = {}
    roles = data.get('roles', [])
    keyboard = [[InlineKeyboardButton(role, callback_data=f'role_{role}')] for role in roles]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text(ui_text.get('start_message'), reply_markup=reply_markup)
    return ROLE

async def role_chosen(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    role = query.data.split('_')[1]
    USER_DATA[user_id]['role'] = role
    faculties = data.get('faculties', [])
    keyboard = [[InlineKeyboardButton(faculty, callback_data=f'faculty_{faculty}')] for faculty in faculties]
    reply_markup = InlineKeyboardMarkup(keyboard)
    current_selection = get_current_selection(USER_DATA[user_id])
    message_text = f"{current_selection}\n{ui_text.get('choose_faculty')}"
    await query.edit_message_text(text=message_text, reply_markup=reply_markup)
    return FACULTY

async def faculty_chosen(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    faculty = query.data.split('_')[1]
    USER_DATA[user_id]['faculty'] = faculty
    departments = data['departments'].get(faculty, [])
    keyboard = [[InlineKeyboardButton(dept, callback_data=f'department_{dept}')] for dept in departments]
    reply_markup = InlineKeyboardMarkup(keyboard)
    current_selection = get_current_selection(USER_DATA[user_id])
    message_text = f"{current_selection}\n{ui_text.get('choose_department')}"
    await query.edit_message_text(text=message_text, reply_markup=reply_markup)
    return DEPARTMENT

async def department_chosen(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    department = query.data.split('_')[1]
    USER_DATA[user_id]['department'] = department
    role = USER_DATA[user_id].get('role')
    popular_docs = data.get('popular_documents', {}).get(role, [])
    all_docs = data.get('documents', {}).get(role, [])
    keyboard = [[InlineKeyboardButton(doc, callback_data=f'document_{doc}')] for doc in popular_docs]
    if all_docs and popular_docs != all_docs:
        keyboard.append([InlineKeyboardButton(ui_text.get('all_documents_button'), callback_data='show_all_documents')])
    reply_markup = InlineKeyboardMarkup(keyboard)
    current_selection = get_current_selection(USER_DATA[user_id])
    message_text = f"{current_selection}\n{ui_text.get('choose_document')}\n{ui_text.get('popular_documents')}"
    await query.edit_message_text(text=message_text, reply_markup=reply_markup)
    return DOCUMENT

async def show_all_documents(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    role = USER_DATA[user_id].get('role')
    all_docs = data.get('documents', {}).get(role, [])
    keyboard = [[InlineKeyboardButton(doc, callback_data=f'doc_id_{i}')] for i, doc in enumerate(all_docs)]
    reply_markup = InlineKeyboardMarkup(keyboard)
    current_selection = get_current_selection(USER_DATA[user_id])
    message_text = f"{current_selection}\n{ui_text.get('choose_document')}"
    context.user_data['all_documents'] = {f'doc_id_{i}': doc for i, doc in enumerate(all_docs)}
    await query.edit_message_text(text=message_text, reply_markup=reply_markup)

async def document_chosen(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    callback_data = query.data
    if callback_data.startswith('doc_id_'):
        document = context.user_data['all_documents'][callback_data]
    else:
        document = callback_data.split('_')[1]
    USER_DATA[user_id]['document'] = document
    await query.edit_message_text(f"{get_current_selection(USER_DATA[user_id])}\n{ui_text.get('enter_full_name')}")
    return FULL_NAME

async def full_name_received(update: Update, context: CallbackContext) -> int:
    user_id = update.effective_user.id
    full_name = update.message.text
    if re.match(FULL_NAME_REGEX, full_name):
        USER_DATA[user_id]['full_name'] = full_name
        document = USER_DATA[user_id]['document']
        fields_to_ask = data.get('document_fields', {}).get(document, [])
        context.user_data['fields_to_ask'] = fields_to_ask[:]
        if fields_to_ask:
            field = fields_to_ask.pop(0)
            await update.message.reply_text(f"{get_current_selection(USER_DATA[user_id])}\n{ui_text.get('enter_your')} {field}:")
            return ADDITIONAL_DATA
        else:
            return await display_confirmation(update, context)
    else:
        await update.message.reply_text(ui_text.get('incorrect_full_name'))
        return FULL_NAME

async def additional_data_received(update: Update, context: CallbackContext) -> int:
    user_id = update.effective_user.id
    field_name = context.user_data['fields_to_ask'][0]
    field_value = update.message.text
    if 'additional_data' not in USER_DATA[user_id]:
        USER_DATA[user_id]['additional_data'] = {}
    USER_DATA[user_id]['additional_data'][field_name] = field_value
    context.user_data['fields_to_ask'].pop(0)
    if context.user_data['fields_to_ask']:
        next_field = context.user_data['fields_to_ask'][0]
        await update.message.reply_text(f"{get_current_selection(USER_DATA[user_id])}\n{ui_text.get('enter_your')} {next_field}:")
        return ADDITIONAL_DATA
    else:
        return await display_confirmation(update, context)

async def display_confirmation(update: Update, context: CallbackContext) -> int:
    user_id = update.effective_user.id
    user_data = USER_DATA[user_id]
    confirmation_text = ui_text.get('confirm_data') + "\n\n"
    for key, value in user_data.items():
        if key not in ['role', 'faculty', 'department', 'document']:
            confirmation_text += f"{key.replace('_', ' ').title()}: {value}\n"
        elif key in ['faculty', 'department']:
            confirmation_text += f"{key.title()}: {value}\n"
        elif key == 'role':
            confirmation_text += f"{key.title()}: {value}\n"
        elif key == 'document':
            confirmation_text += f"{key.title()}: {value}\n"
    if 'additional_data' in user_data:
        confirmation_text += f"\n{ui_text.get('additional_data_label', 'Додаткові дані')}:\n"
        for key, value in user_data['additional_data'].items():
            confirmation_text += f"- {key.replace('_', ' ').title()}: {value}\n"
    confirmation_text += "\nОберіть формат документа:"

    keyboard = [
        [InlineKeyboardButton("DOCX" if context.user_data.get('output_format') != 'docx' else "✅ DOCX", callback_data='format_docx'),
         InlineKeyboardButton("PDF" if context.user_data.get('output_format') != 'pdf' else "✅ PDF", callback_data='format_pdf')],
        [InlineKeyboardButton(ui_text.get('confirm_button'), callback_data='confirm')],
        [InlineKeyboardButton(ui_text.get('change_button'), callback_data='change')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    if update.message:
        await update.message.reply_text(confirmation_text, reply_markup=reply_markup)
    elif update.callback_query:
        query = update.callback_query
        await query.answer()
        await query.edit_message_text(confirmation_text, reply_markup=reply_markup)
    return CONFIRMATION

async def format_chosen(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    await query.answer()
    format_choice = query.data.split('_')[1]  # 'docx' або 'pdf'
    context.user_data['output_format'] = format_choice
    
    # Оновлюємо текст із вибраним форматом
    user_id = update.effective_user.id
    user_data = USER_DATA[user_id]
    confirmation_text = ui_text.get('confirm_data') + "\n\n"
    for key, value in user_data.items():
        if key not in ['role', 'faculty', 'department', 'document']:
            confirmation_text += f"{key.replace('_', ' ').title()}: {value}\n"
        elif key in ['faculty', 'department']:
            confirmation_text += f"{key.title()}: {value}\n"
        elif key == 'role':
            confirmation_text += f"{key.title()}: {value}\n"
        elif key == 'document':
            confirmation_text += f"{key.title()}: {value}\n"
    if 'additional_data' in user_data:
        confirmation_text += f"\n{ui_text.get('additional_data_label', 'Додаткові дані')}:\n"
        for key, value in user_data['additional_data'].items():
            confirmation_text += f"- {key.replace('_', ' ').title()}: {value}\n"
    confirmation_text += f"\nОбраний формат: {format_choice.upper()}"

    # Оновлюємо клавіатуру з позначкою вибору
    keyboard = [
        [InlineKeyboardButton("✅ DOCX" if format_choice == 'docx' else "DOCX", callback_data='format_docx'),
         InlineKeyboardButton("✅ PDF" if format_choice == 'pdf' else "PDF", callback_data='format_pdf')],
        [InlineKeyboardButton(ui_text.get('confirm_button'), callback_data='confirm')],
        [InlineKeyboardButton(ui_text.get('change_button'), callback_data='change')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await query.edit_message_text(confirmation_text, reply_markup=reply_markup)
    return CONFIRMATION

async def confirm_data(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    await query.answer()
    user_id = update.effective_user.id
    user_data = USER_DATA[user_id]
    document_name = user_data.get('document', ui_text.get('the_document', 'документ'))
    user_data_formatted = "\n".join([f"{k.replace('_', ' ').title()}: {v}" for k, v in user_data.items() if k not in ['role', 'faculty', 'department', 'document', 'additional_data']])
    if 'additional_data' in user_data:
        user_data_formatted += f"\n{ui_text.get('additional_data_label', 'Додаткові дані')}:\n"
        user_data_formatted += "\n".join([f"- {k.replace('_', ' ').title()}: {v}" for k, v in user_data['additional_data'].items()])
    output_format = context.user_data.get('output_format', 'docx')
    try:
        file_path = generate_document(user_data, output_format=output_format)
        await query.edit_message_text(
            ui_text.get('document_generation_successful').format(document_name=document_name, user_data=user_data_formatted)
        )
        with open(file_path, 'rb') as f:
            await query.message.reply_document(f, filename=f"{document_name}.{output_format}")
        os.remove(file_path)
    except Exception as e:
        await query.edit_message_text(f"Помилка при генерації документа: {str(e)}")
        return ConversationHandler.END
    del USER_DATA[user_id]
    return ConversationHandler.END

async def change_data(update: Update, context: CallbackContext) -> int:
    query = update.callback_query
    await query.answer()
    first_name = update.effective_user.first_name
    await query.edit_message_text(ui_text.get('data_change_prompt'))
    return ConversationHandler.END

async def cancel(update: Update, context: CallbackContext) -> int:
    user = update.effective_user
    await update.message.reply_text(ui_text.get('cancel_message').format(first_name=user.first_name))
    return ConversationHandler.END

def main() -> None:
    bot_token = os.getenv("BOT_TOKEN")
    if not bot_token:
        print("Error: BOT_TOKEN not found in .env file.")
        return
    application = ApplicationBuilder().token(bot_token).build()
    conv_handler = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            ROLE: [CallbackQueryHandler(role_chosen, pattern='^role_')],
            FACULTY: [CallbackQueryHandler(faculty_chosen, pattern='^faculty_')],
            DEPARTMENT: [CallbackQueryHandler(department_chosen, pattern='^department_')],
            DOCUMENT: [
                CallbackQueryHandler(document_chosen, pattern=lambda data: data.startswith('document_') or data.startswith('doc_id_')),
                CallbackQueryHandler(show_all_documents, pattern='^show_all_documents$')
            ],
            FULL_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, full_name_received)],
            ADDITIONAL_DATA: [MessageHandler(filters.TEXT & ~filters.COMMAND, additional_data_received)],
            CONFIRMATION: [
                CallbackQueryHandler(confirm_data, pattern='^confirm$'),
                CallbackQueryHandler(change_data, pattern='^change$'),
                CallbackQueryHandler(format_chosen, pattern='^format_'),
            ],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )
    application.add_handler(conv_handler)
    application.run_polling()

if __name__ == '__main__':
    main()