import os
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

class DocumentGenerator:
    def generate(self, user_data, output_format='docx'):
        file_name = f"{user_data['full_name']}_{user_data['document']}"
        if output_format.lower() == 'docx':
            return self._generate_docx(user_data, file_name)
        elif output_format.lower() == 'pdf':
            return self._generate_pdf(user_data, file_name)
        else:
            raise ValueError("Непідтримуваний формат. Використовуйте 'docx' або 'pdf'.")

    def _generate_docx(self, user_data, file_name):
        doc = Document()
        doc.add_paragraph(f"Document: {user_data['document']}")
        doc.add_paragraph(f"Full Name: {user_data['full_name']}")
        if 'phone_number' in user_data:
            doc.add_paragraph(f"Phone Number: {user_data['phone_number']}")
        if 'role' in user_data:
            doc.add_paragraph(f"Role: {user_data['role']}")
        if 'faculty' in user_data:
            doc.add_paragraph(f"Faculty: {user_data['faculty']}")
        if 'department' in user_data:
            doc.add_paragraph(f"Department: {user_data['department']}")
        if 'education_degree' in user_data:
            doc.add_paragraph(f"Education Degree: {user_data['education_degree']}")
        if 'speciality' in user_data:
            doc.add_paragraph(f"Speciality: {user_data['speciality']}")
        if 'course' in user_data:
            doc.add_paragraph(f"Course: {user_data['course']}")
        if 'position' in user_data:
            doc.add_paragraph(f"Position: {user_data['position']}")
        if 'additional_data' in user_data:
            for key, value in user_data['additional_data'].items():
                doc.add_paragraph(f"{key.title()}: {value}")
        file_path = f"{file_name}.docx"
        doc.save(file_path)
        return file_path

    def _generate_pdf(self, user_data, file_name):
        file_path = f"{file_name}.pdf"
        c = canvas.Canvas(file_path, pagesize=letter)
        width, height = letter
        y_position = height - 50
        c.drawString(100, y_position, f"Document: {user_data['document']}")
        y_position -= 20
        c.drawString(100, y_position, f"Full Name: {user_data['full_name']}")
        y_position -= 20
        if 'phone_number' in user_data:
            c.drawString(100, y_position, f"Phone Number: {user_data['phone_number']}")
            y_position -= 20
        if 'role' in user_data:
            c.drawString(100, y_position, f"Role: {user_data['role']}")
            y_position -= 20
        if 'faculty' in user_data:
            c.drawString(100, y_position, f"Faculty: {user_data['faculty']}")
            y_position -= 20
        if 'department' in user_data:
            c.drawString(100, y_position, f"Department: {user_data['department']}")
            y_position -= 20
        if 'education_degree' in user_data:
            c.drawString(100, y_position, f"Education Degree: {user_data['education_degree']}")
            y_position -= 20
        if 'speciality' in user_data:
            c.drawString(100, y_position, f"Speciality: {user_data['speciality']}")
            y_position -= 20
        if 'course' in user_data:
            c.drawString(100, y_position, f"Course: {user_data['course']}")
            y_position -= 20
        if 'position' in user_data:
            c.drawString(100, y_position, f"Position: {user_data['position']}")
            y_position -= 20
        if 'additional_data' in user_data:
            for key, value in user_data['additional_data'].items():
                c.drawString(100, y_position, f"{key.title()}: {value}")
                y_position -= 20
        c.showPage()
        c.save()
        return file_path
# import os
# from docx import Document
# from reportlab.lib.pagesizes import letter
# from reportlab.pdfgen import canvas

# class DocumentGenerator:
#     def generate(self, user_data, output_format='docx'):
#         file_name = f"{user_data['full_name']}_{user_data['document']}"
#         if output_format.lower() == 'docx':
#             return self._generate_docx(user_data, file_name)
#         elif output_format.lower() == 'pdf':
#             return self._generate_pdf(user_data, file_name)
#         else:
#             raise ValueError("Непідтримуваний формат. Використовуйте 'docx' або 'pdf'.")

#     def _generate_docx(self, user_data, file_name):
#         doc = Document()
#         doc.add_paragraph(f"Document: {user_data['document']}")
#         doc.add_paragraph(f"Full Name: {user_data['full_name']}")
#         if 'phone_number' in user_data:
#             doc.add_paragraph(f"Phone Number: {user_data['phone_number']}")
#         if 'additional_data' in user_data:
#             for key, value in user_data['additional_data'].items():
#                 doc.add_paragraph(f"{key.title()}: {value}")
#         file_path = f"{file_name}.docx"
#         doc.save(file_path)
#         return file_path

#     def _generate_pdf(self, user_data, file_name):
#         file_path = f"{file_name}.pdf"
#         c = canvas.Canvas(file_path, pagesize=letter)
#         width, height = letter
#         y_position = height - 50
#         c.drawString(100, y_position, f"Document: {user_data['document']}")
#         y_position -= 20
#         c.drawString(100, y_position, f"Full Name: {user_data['full_name']}")
#         y_position -= 20
#         if 'phone_number' in user_data:
#             c.drawString(100, y_position, f"Phone Number: {user_data['phone_number']}")
#             y_position -= 20
#         if 'additional_data' in user_data:
#             for key, value in user_data['additional_data'].items():
#                 c.drawString(100, y_position, f"{key.title()}: {value}")
#                 y_position -= 20
#         c.showPage()
#         c.save()
#         return file_path